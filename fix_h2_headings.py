"""
Fix empty H2 paragraphs: pull bold heading text from the next Normal paragraph
into the empty H2, leaving remaining body text in the Normal paragraph.
"""
import docx
from copy import deepcopy
from lxml import etree

FILEPATH = r'c:\Users\bchance\Dropbox\My Documents\Classes\Stat 301 - Win 26\lectures\finalreview.docx'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def qn(tag):
    return f'{{{W_NS}}}{tag}'

doc = docx.Document(FILEPATH)

fixed = 0
i = 0
while i < len(doc.paragraphs):
    p = doc.paragraphs[i]
    # Find empty H2 paragraphs
    if p.style.name == 'Heading 2' and not p.text.strip():
        # Check if next paragraph exists and has bold-prefixed runs
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            if next_p.runs and next_p.runs[0].bold:
                # Find split point: consecutive bold runs from start
                split_at = 0
                for j, run in enumerate(next_p.runs):
                    if run.bold:
                        split_at = j + 1
                    else:
                        break

                # Move bold runs from next_p into the empty H2
                p_elem = p._element
                next_elem = next_p._element

                # Get the run elements to move
                all_runs = [r._element for r in next_p.runs]
                heading_runs = all_runs[:split_at]
                
                # Move heading runs into the H2 paragraph
                for run_elem in heading_runs:
                    p_elem.append(run_elem)  # This moves (not copies) the element

                # If next paragraph is now empty, remove it
                remaining_text = next_p.text.strip()
                if not remaining_text:
                    next_elem.getparent().remove(next_elem)

                fixed += 1
                heading_text = p.text.strip()[:60]
                print(f"  Fixed P{i}: '{heading_text}'")
    i += 1

print(f"\nFixed {fixed} empty H2 paragraphs")
doc.save(FILEPATH)
print(f"Saved to: {FILEPATH}")
