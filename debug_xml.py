import docx
from lxml import etree
import re

doc = docx.Document(r'c:\Users\bchance\Dropbox\My Documents\Classes\Stat 301 - Win 26\lectures\finalreview.docx')

# Look at P16 (body paragraph after empty "Exam Format" H2)
p = doc.paragraphs[16]
xml = etree.tostring(p._element, pretty_print=True).decode()
# Remove namespace declarations for readability
clean = re.sub(r'\s+xmlns:\w+="[^"]*"', '', xml)
print("P16 XML (first 3000 chars):")
print(clean[:3000])

print("\n\n")
# Also look at P3 (body after empty "Due Tuesday" H2)  
p3 = doc.paragraphs[3]
xml3 = etree.tostring(p3._element, pretty_print=True).decode()
clean3 = re.sub(r'\s+xmlns:\w+="[^"]*"', '', xml3)
print("P3 XML (first 3000 chars):")
print(clean3[:3000])

# Also check: what child element tags does P3 have?
print("\n\nP3 direct children tags:")
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for child in p3._element:
    tag = child.tag.replace('{' + W_NS + '}', 'w:')
    text = child.text if child.text else ''
    print(f"  {tag}: text={text[:30]!r}")
