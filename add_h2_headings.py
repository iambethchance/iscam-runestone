"""
Convert bold-prefixed paragraphs in a Word document to Heading 2 style.
For mixed paragraphs (bold heading + body text), splits them into
a Heading 2 paragraph + a Normal body paragraph.
"""
import docx
from docx.shared import Pt
from copy import deepcopy
from lxml import etree

FILEPATH = r'c:\Users\bchance\Dropbox\My Documents\Classes\Stat 301 - Win 26\lectures\finalreview.docx'

doc = docx.Document(FILEPATH)

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def qn(tag):
    """Create a qualified name in the w: namespace."""
    return f'{{{W_NS}}}{tag}'

# Identify paragraphs to process: bold first run, has text, not already Heading 1
to_process = []
for i, p in enumerate(doc.paragraphs):
    if p.runs and p.runs[0].bold and p.text.strip() and p.style.name != 'Heading 1':
        to_process.append(i)

print(f"Found {len(to_process)} paragraphs to convert to Heading 2")

# Process in reverse order so inserting new paragraphs doesn't shift indices
for idx in reversed(to_process):
    p = doc.paragraphs[idx]
    p_elem = p._element

    # Find split point: count consecutive bold runs from the start
    split_at = 0
    for j, run in enumerate(p.runs):
        if run.bold:
            split_at = j + 1
        else:
            break

    # Check if there are non-bold runs with actual text content
    if split_at < len(p.runs):
        body_text = ''.join(r.text for r in p.runs[split_at:]).strip()
    else:
        body_text = ''

    has_body = bool(body_text)

    if has_body:
        # Collect the XML elements for heading runs (to identify which to keep/remove)
        heading_run_ids = set()
        for j in range(split_at):
            heading_run_ids.add(id(p.runs[j]._element))

        # Get all direct children that are not paragraph properties
        children = [c for c in p_elem if c.tag != qn('pPr')]

        # Split children into heading group and body group
        heading_children = []
        body_children = []
        in_body = False

        for child in children:
            if not in_body and child.tag == qn('r') and id(child) not in heading_run_ids:
                in_body = True
            if in_body:
                body_children.append(child)
            else:
                heading_children.append(child)

        if body_children:
            # Create a new paragraph element for the body content
            new_p_elem = deepcopy(p_elem)

            # Clear all content from the new paragraph (keep only pPr)
            for child in list(new_p_elem):
                if child.tag != qn('pPr'):
                    new_p_elem.remove(child)

            # Add body children (deep copies) to new paragraph
            for child in body_children:
                new_p_elem.append(deepcopy(child))

            # Remove body children from original paragraph
            for child in body_children:
                p_elem.remove(child)

            # Set new paragraph style to Normal
            pPr = new_p_elem.find(qn('pPr'))
            if pPr is None:
                pPr = etree.SubElement(new_p_elem, qn('pPr'))
                # Insert pPr as first child
                new_p_elem.insert(0, pPr)
            pStyle = pPr.find(qn('pStyle'))
            if pStyle is not None:
                pStyle.set(qn('val'), 'Normal')
            else:
                pStyle = etree.SubElement(pPr, qn('pStyle'))
                pStyle.set(qn('val'), 'Normal')

            # Insert new body paragraph after the original heading paragraph
            p_elem.addnext(new_p_elem)

    # Change original paragraph style to Heading 2
    p.style = doc.styles['Heading 2']

    # Set font properties for all runs in the heading: bold, Arial
    for run in p.runs:
        run.bold = True
        run.font.name = 'Arial'

    heading_text = p.text.strip()[:60]
    print(f"  P{idx}: -> H2: {heading_text!r}")

# Save the modified document
doc.save(FILEPATH)
print(f"\nSaved to: {FILEPATH}")
